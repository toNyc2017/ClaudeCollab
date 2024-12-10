from docx import Document
from playwright.sync_api import sync_playwright
import time
from datetime import datetime
import os
import json
import sys
from pathlib import Path

class EmailExtractor:
    def __init__(self, email, password):
        self.email = email
        self.password = password
        self.progress_file = '/tmp/extraction_progress.json'
        self.document = None
        self.text_content = []
        self.load_progress()

    def load_progress(self):
        try:
            if os.path.exists(self.progress_file):
                with open(self.progress_file, 'r') as f:
                    self.progress = json.load(f)
                print(f"Loaded previous progress: {self.progress['processed_count']} emails processed")
            else:
                self.progress = {'processed_count': 0, 'last_file': None}
        except Exception as e:
            print(f"Error loading progress: {e}")
            self.progress = {'processed_count': 0, 'last_file': None}

    def save_progress(self, filename=None):
        try:
            if filename:
                self.progress['last_file'] = filename
            with open(self.progress_file, 'w') as f:
                json.dump(self.progress, f)
        except Exception as e:
            print(f"Error saving progress: {e}")

    def retry_operation(self, operation, max_retries=3, delay=2):
        for attempt in range(max_retries):
            try:
                return operation()
            except Exception as e:
                if attempt == max_retries - 1:
                    raise e
                print(f"Attempt {attempt + 1} failed: {e}")
                print(f"Retrying in {delay} seconds...")
                time.sleep(delay)
                delay *= 2

    def wait_for_load(self, page, timeout=30):
        print("Waiting for page to load completely...")
        deadline = time.time() + timeout
        while time.time() < deadline:
            try:
                page.wait_for_load_state("networkidle", timeout=5000)
                page.wait_for_load_state("domcontentloaded", timeout=5000)
                print("Page loaded successfully")
                return True
            except Exception as e:
                print(f"Still waiting for page to load... ({str(e)})")
                time.sleep(2)
        raise Exception("Page failed to load completely")

    def login_to_outlook(self, page):
        def _login():
            print("\nLogging in to Outlook...")
            page.goto("https://outlook.office365.com/", wait_until="networkidle")
            self.wait_for_load(page)
            time.sleep(2)

            # Handle sign in button
            sign_in_elements = page.query_selector_all("text=Sign in")
            for element in sign_in_elements:
                if element.is_visible():
                    print("Clicking sign in button...")
                    element.click()
                    break

            # Handle email input
            print("Entering email...")
            email_input = page.wait_for_selector("input[type='email']", timeout=10000)
            email_input.fill(self.email)
            time.sleep(1)
            page.keyboard.press("Enter")

            # Handle password input
            print("Entering password...")
            password_input = page.wait_for_selector("input[type='password']", timeout=10000)
            password_input.fill(self.password)
            time.sleep(1)
            page.keyboard.press("Enter")

            # Handle 2FA if needed
            try:
                auth_code_element = page.wait_for_selector("text=Enter code", timeout=10000)
                if auth_code_element:
                    print("\n*** 2FA Required ***")
                    print("Please enter the authentication code manually")
                    input("Press Enter after entering the code...")
            except:
                print("No 2FA prompt detected")

            # Wait for main interface
            print("Waiting for Outlook to load...")
            self.wait_for_load(page)
            time.sleep(5)
            return True

        return self.retry_operation(_login)

    def find_search_box(self, page):
        print("\nLooking for search box...")
        search_selectors = [
            "input[aria-label='Search']",
            "input[placeholder='Search']",
            "[role='searchbox']",
            "input[type='text']",
            "input[type='search']",
            "[aria-label*='search' i]",
            "[placeholder*='search' i]",
            "input"
        ]
        
        for selector in search_selectors:
            try:
                print(f"Trying selector: {selector}")
                elements = page.query_selector_all(selector)
                for element in elements:
                    if element.is_visible():
                        try:
                            box_text = element.get_attribute('aria-label') or element.get_attribute('placeholder') or ''
                            print(f"Found potential search box: {box_text}")
                            return element
                        except:
                            continue
            except Exception as e:
                print(f"Error with selector {selector}: {e}")
        
        # Take screenshot for debugging
        page.screenshot(path="/tmp/outlook_debug.png")
        print("Debug screenshot saved as /tmp/outlook_debug.png")
        
        raise Exception("Could not find search bar")

    def navigate_to_email_list(self, page):
        print("Attempting to return to email list...")
        
        try:
            # Try browser back button first
            page.go_back()
            time.sleep(2)
            
            # Check if we're back at the list
            if page.query_selector('div[role="listitem"]'):
                print("Successfully returned to list using browser back")
                return
            
            # Try clicking "Mail" in the left sidebar
            mail_link = page.query_selector("a:has-text('Mail')")
            if mail_link:
                mail_link.click()
                time.sleep(2)
                if page.query_selector('div[role="listitem"]'):
                    print("Successfully returned to list using Mail link")
                    return
            
            # Try searching again
            search_bar = self.find_search_box(page)
            if search_bar:
                search_bar.click()
                time.sleep(1)
                page.keyboard.press("Control+a")
                page.keyboard.press("Backspace")
                time.sleep(1)
                search_bar.fill("from:Lynn Gadue")
                search_bar.press("Enter")
                time.sleep(3)
                if page.query_selector('div[role="listitem"]'):
                    print("Successfully returned to list using search")
                    return
            
            raise Exception("Could not return to email list")
        
        except Exception as e:
            print(f"Error navigating to email list: {e}")
            page.screenshot(path="/tmp/outlook_navigation_error.png")
            raise

    def init_document(self):
        if not self.document:
            self.document = Document()
            self.document.add_heading('Emails from Lynn Gadue', 0)
            self.document.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            self.document.add_paragraph('Last updated: [will be updated on save]')
            
            # Initialize text content
            self.text_content = [
                'Emails from Lynn Gadue',
                f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}',
                ''
            ]

    def save_document(self):
        if self.document:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Update timestamps
            update_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            for para in self.document.paragraphs:
                if para.text.startswith('Last updated:'):
                    para.text = f'Last updated: {update_time}'
                    break

            # Save Word document
            docx_filename = f'/tmp/lynn_gadue_emails_{timestamp}.docx'
            self.document.save(docx_filename)
            print(f'\nSaved Word document to {docx_filename}')
            
            # Save text version
            text_filename = f'/tmp/lynn_gadue_emails_{timestamp}.txt'
            with open(text_filename, 'w', encoding='utf-8') as f:
                f.write('\n'.join(self.text_content))
            print(f'Saved text version to {text_filename}')
            
            return docx_filename
        return None

    def extract_email_content(self, page):
        try:
            print("Extracting email content...")
            content_selectors = [
                'div[role="document"]',
                'div[role="main"]',
                '[role="region"]',
                '.message-body',
                '.email-content'
            ]
            
            for selector in content_selectors:
                content = page.query_selector(selector)
                if content and content.is_visible():
                    return content.inner_text()
            
            # If no content found, take debug screenshot
            page.screenshot(path="/tmp/outlook_content_error.png")
            raise Exception("Could not find email content")
            
        except Exception as e:
            print(f"Error extracting content: {e}")
            raise

    def extract_emails(self, page):
        try:
            self.init_document()

            # Wait for page load
            print("\nPreparing to search for emails...")
            self.wait_for_load(page)
            time.sleep(3)

            # Find and use search box
            search_bar = self.retry_operation(lambda: self.find_search_box(page))
            
            print("Clearing search box...")
            search_bar.click()
            time.sleep(1)
            page.keyboard.press("Control+a")
            page.keyboard.press("Backspace")
            time.sleep(1)
            
            print("Entering search term...")
            search_bar.fill("from:Lynn Gadue")
            search_bar.press("Enter")

            print("Waiting for search results...")
            self.wait_for_load(page)
            time.sleep(3)

            # Process emails
            email_items = page.query_selector_all('div[role="listitem"]')
            total_emails = len(email_items)
            print(f'\nFound {total_emails} emails from Lynn Gadue')

            if total_emails == 0:
                # Take debug screenshot
                page.screenshot(path="/tmp/outlook_no_results.png")
                print("Debug screenshot saved as /tmp/outlook_no_results.png")
                raise Exception("No emails found")

            # Skip already processed emails
            start_index = self.progress['processed_count']
            if start_index > 0:
                print(f"Resuming from email {start_index + 1}")

            for i, item in enumerate(email_items[start_index:], start=start_index + 1):
                try:
                    print(f'\nProcessing email {i}/{total_emails}')
                    
                    # Extract metadata
                    subject_element = item.query_selector('div[role="heading"]')
                    date_element = item.query_selector('div[title*="/"]')
                    
                    subject_text = subject_element.inner_text() if subject_element else "No subject"
                    date_text = date_element.get_attribute('title') if date_element else "Date unknown"
                    
                    print(f'Subject: {subject_text}')
                    print(f'Date: {date_text}')

                    # Open and extract email
                    item.click()
                    time.sleep(2)
                    self.wait_for_load(page)

                    content_text = self.extract_email_content(page)
                    if content_text:
                        # Add to Word document
                        self.document.add_heading(f'Email {i}', level=1)
                        self.document.add_paragraph(f'Date: {date_text}')
                        self.document.add_paragraph(f'Subject: {subject_text}')
                        self.document.add_paragraph('Content:')
                        self.document.add_paragraph(content_text)
                        self.document.add_paragraph('---')
                        
                        # Add to text content
                        self.text_content.extend([
                            f'\nEmail {i}',
                            f'Date: {date_text}',
                            f'Subject: {subject_text}',
                            'Content:',
                            content_text,
                            '---\n'
                        ])
                        
                        print(f'Added content ({len(content_text)} chars)')

                        # Save progress every 2 emails
                        if i % 2 == 0:
                            filename = self.save_document()
                            self.progress['processed_count'] = i
                            self.save_progress(filename)

                    # Return to email list
                    self.navigate_to_email_list(page)
                    self.wait_for_load(page)

                except Exception as e:
                    print(f"Error processing email {i}: {e}")
                    # Save progress before continuing
                    self.save_document()
                    self.progress['processed_count'] = i - 1
                    self.save_progress()
                    continue

            # Final save
            final_filename = self.save_document()
            self.progress['processed_count'] = total_emails
            self.save_progress(final_filename)
            return final_filename

        except Exception as e:
            print(f"Error during extraction: {e}")
            return self.save_document()

def main():
    email = "tolds@3clife.info"
    password = "annuiTy2024!"
    
    extractor = EmailExtractor(email, password)

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=False)
        context = browser.new_context(viewport={'width': 1280, 'height': 800})
        page = context.new_page()
        
        try:
            if extractor.login_to_outlook(page):
                print("\nLogin successful!")
                filename = extractor.extract_emails(page)
                if filename:
                    print("\nExtraction completed!")
                    print(f"Final document saved as: {filename}")
                else:
                    print("\nExtraction failed!")
            else:
                print("\nLogin failed!")
        except Exception as e:
            print(f"\nError: {e}")
        finally:
            print("\nClosing browser...")
            context.close()
            browser.close()

if __name__ == "__main__":
    main()