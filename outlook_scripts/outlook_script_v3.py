from docx import Document
from playwright.sync_api import sync_playwright
import time
from datetime import datetime
import os
import json
import sys
from pathlib import Path

class EmailExtractor:
    def __init__(self, email, password):
        self.email = email
        self.password = password
        self.progress_file = '/tmp/extraction_progress.json'
        self.document = None
        self.load_progress()

    def load_progress(self):
        try:
            if os.path.exists(self.progress_file):
                with open(self.progress_file, 'r') as f:
                    self.progress = json.load(f)
                print(f"Loaded previous progress: {self.progress['processed_count']} emails processed")
            else:
                self.progress = {'processed_count': 0, 'last_file': None}
        except Exception as e:
            print(f"Error loading progress: {e}")
            self.progress = {'processed_count': 0, 'last_file': None}

    def save_progress(self, filename=None):
        try:
            if filename:
                self.progress['last_file'] = filename
            with open(self.progress_file, 'w') as f:
                json.dump(self.progress, f)
        except Exception as e:
            print(f"Error saving progress: {e}")

    def retry_operation(self, operation, max_retries=3, delay=2):
        for attempt in range(max_retries):
            try:
                return operation()
            except Exception as e:
                if attempt == max_retries - 1:
                    raise e
                print(f"Attempt {attempt + 1} failed: {e}")
                print(f"Retrying in {delay} seconds...")
                time.sleep(delay)
                delay *= 2

    def login_to_outlook(self, page):
        def _login():
            print("\nLogging in to Outlook...")
            page.goto("https://outlook.office365.com/", wait_until="networkidle")
            page.wait_for_load_state("domcontentloaded")
            time.sleep(2)

            # Handle sign in button
            sign_in_elements = page.query_selector_all("text=Sign in")
            for element in sign_in_elements:
                if element.is_visible():
                    print("Clicking sign in button...")
                    element.click()
                    break

            # Handle email input
            print("Entering email...")
            email_input = page.wait_for_selector("input[type='email']", timeout=10000)
            email_input.fill(self.email)
            time.sleep(1)
            page.keyboard.press("Enter")

            # Handle password input
            print("Entering password...")
            password_input = page.wait_for_selector("input[type='password']", timeout=10000)
            password_input.fill(self.password)
            time.sleep(1)
            page.keyboard.press("Enter")

            # Handle 2FA if needed
            try:
                auth_code_element = page.wait_for_selector("text=Enter code", timeout=10000)
                if auth_code_element:
                    print("\n*** 2FA Required ***")
                    print("Please enter the authentication code manually")
                    input("Press Enter after entering the code...")
            except:
                print("No 2FA prompt detected")

            # Wait for main interface
            print("Waiting for Outlook to load...")
            time.sleep(5)
            return True

        return self.retry_operation(_login)

    def init_document(self):
        if not self.document:
            self.document = Document()
            self.document.add_heading('Emails from Lynn Gadue', 0)
            self.document.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            self.document.add_paragraph('Last updated: [will be updated on save]')

    def save_document(self):
        if self.document:
            # Update the last updated timestamp
            for para in self.document.paragraphs:
                if para.text.startswith('Last updated:'):
                    para.text = f'Last updated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
                    break

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f'/tmp/lynn_gadue_emails_{timestamp}.docx'
            self.document.save(filename)
            print(f'\nSaved document to {filename}')
            return filename
        return None

    def extract_emails(self, page):
        try:
            self.init_document()

            # Search for emails
            print("\nSearching for Lynn's emails...")
            page.wait_for_load_state("networkidle")
            time.sleep(3)

            # Find search box
            search_selectors = [
                "input[aria-label='Search']",
                "input[placeholder='Search']",
                "[role='searchbox']",
                "input[type='text']"
            ]

            search_bar = None
            for selector in search_selectors:
                try:
                    search_bar = page.wait_for_selector(selector, timeout=5000)
                    if search_bar and search_bar.is_visible():
                        print(f"Found search bar using: {selector}")
                        break
                except:
                    continue

            if not search_bar:
                raise Exception("Could not find search bar")

            # Perform search
            search_bar.click()
            time.sleep(1)
            page.keyboard.press("Control+a")
            page.keyboard.press("Backspace")
            time.sleep(1)
            search_bar.fill("from:Lynn Gadue")
            search_bar.press("Enter")

            print("Waiting for search results...")
            time.sleep(3)
            page.wait_for_load_state("networkidle")

            # Process emails
            email_items = page.query_selector_all('div[role="listitem"]')
            total_emails = len(email_items)
            print(f'\nFound {total_emails} emails from Lynn Gadue')

            # Skip already processed emails
            start_index = self.progress['processed_count']
            if start_index > 0:
                print(f"Resuming from email {start_index + 1}")

            for i, item in enumerate(email_items[start_index:], start=start_index + 1):
                try:
                    print(f'\nProcessing email {i}/{total_emails}')
                    
                    # Extract metadata
                    subject_element = item.query_selector('div[role="heading"]')
                    date_element = item.query_selector('div[title*="/"]')
                    
                    subject_text = subject_element.inner_text() if subject_element else "No subject"
                    date_text = date_element.get_attribute('title') if date_element else "Date unknown"
                    
                    print(f'Subject: {subject_text}')
                    print(f'Date: {date_text}')

                    # Open and extract email
                    item.click()
                    time.sleep(2)

                    content = page.query_selector('div[role="document"]')
                    if content:
                        content_text = content.inner_text()
                        
                        # Add to document
                        self.document.add_heading(f'Email {i}', level=1)
                        self.document.add_paragraph(f'Date: {date_text}')
                        self.document.add_paragraph(f'Subject: {subject_text}')
                        self.document.add_paragraph('Content:')
                        self.document.add_paragraph(content_text)
                        self.document.add_paragraph('---')
                        
                        print(f'Added content ({len(content_text)} chars)')

                        # Save progress every 5 emails
                        if i % 5 == 0:
                            filename = self.save_document()
                            self.progress['processed_count'] = i
                            self.save_progress(filename)

                    # Return to email list
                    page.go_back()
                    time.sleep(2)
                    page.wait_for_selector('div[role="listitem"]', timeout=10000)

                except Exception as e:
                    print(f"Error processing email {i}: {e}")
                    # Save progress before continuing
                    self.save_document()
                    self.progress['processed_count'] = i - 1
                    self.save_progress()
                    continue

            # Final save
            final_filename = self.save_document()
            self.progress['processed_count'] = total_emails
            self.save_progress(final_filename)
            return final_filename

        except Exception as e:
            print(f"Error during extraction: {e}")
            return self.save_document()

def main():
    email = "tolds@3clife.info"
    password = "annuiTy2024!"
    
    extractor = EmailExtractor(email, password)

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=False)
        context = browser.new_context()
        page = context.new_page()
        
        try:
            if extractor.login_to_outlook(page):
                print("\nLogin successful!")
                filename = extractor.extract_emails(page)
                if filename:
                    print("\nExtraction completed!")
                    print(f"Final document saved as: {filename}")
                else:
                    print("\nExtraction failed!")
            else:
                print("\nLogin failed!")
        except Exception as e:
            print(f"\nError: {e}")
        finally:
            print("\nClosing browser...")
            context.close()
            browser.close()

if __name__ == "__main__":
    main()