from docx import Document
from playwright.sync_api import sync_playwright
import time
from datetime import datetime

def login_to_outlook(page, email, password):
    try:
        print("Navigating to Outlook...")
        page.goto("https://outlook.office365.com/", wait_until="networkidle")
        page.wait_for_load_state("domcontentloaded")
        time.sleep(2)

        # Look for and click any visible "Sign in" button
        print("Looking for 'Sign in' button...")
        sign_in_elements = page.query_selector_all("text=Sign in")
        for element in sign_in_elements:
            if element.is_visible():
                print("Found visible 'Sign in' button, clicking...")
                element.click()
                break

        # Wait for the email input field
        print("Waiting for email input...")
        email_input = page.wait_for_selector("input[type='email']", timeout=10000)
        if email_input:
            print("Found email input, filling...")
            email_input.fill(email)
            time.sleep(1)
            page.keyboard.press("Enter")

        # Wait for the password input field
        print("Waiting for password input...")
        password_input = page.wait_for_selector("input[type='password']", timeout=10000)
        if password_input:
            print("Found password input, filling...")
            password_input.fill(password)
            time.sleep(1)
            page.keyboard.press("Enter")

        # Handle authentication code if needed
        try:
            auth_code_element = page.wait_for_selector("text=Enter code", timeout=10000)
            if auth_code_element:
                # Center the element for better visibility
                auth_code_element.scroll_into_view_if_needed()
                page.evaluate("""element => {
                    element.scrollIntoView({
                        behavior: 'smooth',
                        block: 'center',
                        inline: 'center'
                    });
                }""", auth_code_element)
                
                print("\n*** Authentication code page detected ***")
                print("Please enter the code in your authenticator app")
                input("Press Enter once you've entered the code...")
        except Exception as e:
            print("No authentication code page found, continuing...")

        # Wait for the main interface to load
        print("Waiting for Outlook to load...")
        time.sleep(5)
        return True

    except Exception as e:
        print(f"Login error: {str(e)}")
        return False

def extract_emails(page):
    try:
        document = Document()
        document.add_heading('Emails from Lynn Gadue', 0)
        document.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

        # Search for Lynn's emails
        print("\nSearching for Lynn Gadue's emails...")
        
        # Wait for page to be fully loaded
        page.wait_for_load_state("networkidle")
        time.sleep(3)
        
        # Try different search box selectors
        search_selectors = [
            "input[aria-label='Search']",
            "input[placeholder='Search']",
            "[role='searchbox']",
            "input[type='text']"
        ]
        
        search_bar = None
        for selector in search_selectors:
            try:
                search_bar = page.wait_for_selector(selector, timeout=5000)
                if search_bar and search_bar.is_visible():
                    print(f"Found search bar using selector: {selector}")
                    break
            except:
                continue

        if not search_bar:
            raise Exception("Could not find the search bar")

        # Clear any existing search and perform new search
        search_bar.click()
        time.sleep(1)
        page.keyboard.press("Control+a")  # Select all existing text
        page.keyboard.press("Backspace")  # Clear the field
        time.sleep(1)
        search_bar.fill("from:Lynn Gadue")
        search_bar.press("Enter")
        
        # Wait for search results
        print("Waiting for search results...")
        time.sleep(3)
        page.wait_for_load_state("networkidle")
        
        # Find all email items
        email_items = page.query_selector_all('div[role="listitem"]')
        print(f'\nFound {len(email_items)} emails from Lynn Gadue')
        
        for i, item in enumerate(email_items, 1):
            try:
                # Extract subject and date before clicking
                subject_element = item.query_selector('div[role="heading"]')
                date_element = item.query_selector('div[title*="/"]')  # Look for element with date in title
                
                subject_text = subject_element.inner_text() if subject_element else "No subject"
                date_text = date_element.get_attribute('title') if date_element else "Date unknown"
                
                print(f'\nProcessing email {i}/{len(email_items)}:')
                print(f'Subject: {subject_text}')
                print(f'Date: {date_text}')
                
                # Click to open email
                item.click()
                time.sleep(2)
                
                # Get email content
                content = page.query_selector('div[role="document"]')
                if content:
                    content_text = content.inner_text()
                    
                    # Add to document with date
                    document.add_heading(f'Email {i}', level=1)
                    document.add_paragraph(f'Date: {date_text}')
                    document.add_paragraph(f'Subject: {subject_text}')
                    document.add_paragraph('Content:')
                    document.add_paragraph(content_text)
                    document.add_paragraph('---') # Add separator between emails
                    
                    print(f'Added email content ({len(content_text)} characters)')
                
                # Go back to list
                page.go_back()
                time.sleep(2)
                
                # Wait for the list to be interactive again
                page.wait_for_selector('div[role="listitem"]', timeout=10000)
                
            except Exception as e:
                print(f"Error processing email {i}: {str(e)}")
                continue
        
        # Save document with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f'/tmp/lynn_gadue_emails_{timestamp}.docx'
        document.save(filename)
        print(f'\nSaved emails to {filename}')
        return filename

    except Exception as e:
        print(f"Error extracting emails: {str(e)}")
        return None

if __name__ == "__main__":
    email = "tolds@3clife.info"
    password = "annuiTy2024!"

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=False)
        page = browser.new_page()
        
        try:
            if login_to_outlook(page, email, password):
                print("\nLogin successful!")
                filename = extract_emails(page)
                if filename:
                    print("\nProcess completed successfully!")
                    print(f"Please check {filename} for the extracted emails.")
                else:
                    print("\nFailed to extract emails.")
            else:
                print("\nLogin failed!")
        except Exception as e:
            print(f"\nAn error occurred: {str(e)}")
        
        input("\nPress Enter to close the browser...")